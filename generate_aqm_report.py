#!/usr/bin/env python3
"""
generate_aqm_report.py — Alert Quality Management Report Generator

Generates a professional AQM analysis report from raw NrAiIncident data.
Produces two outputs:
  1. A formatted .docx report with all data tables and template narrative
  2. A structured prompt file (.md) for Claude to generate polished recommendations

Usage:
  python3 generate_aqm_report.py --csv incidents.csv --account "ACME Corp"
  python3 generate_aqm_report.py --csv incidents.csv --account "ACME Corp" --analyst "Jane Doe"

Requirements: pip install pandas python-docx
"""
import argparse, re, sys
from datetime import datetime
from pathlib import Path
import pandas as pd, numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# === CONFIG ===
HEADER_BG = "00AC69"
ROW_ALT_BG = "E8F8F0"
WHITE_BG = "FFFFFF"
PRIMARY = RGBColor(0x00,0xAC,0x69)
SECONDARY = RGBColor(0x33,0xBD,0x87)
DARK = RGBColor(0x1E,0x29,0x3B)
GRAY = RGBColor(0x64,0x74,0x8B)
RED = RGBColor(0xDC,0x26,0x26)
AMBER = RGBColor(0xD9,0x77,0x06)
FONT = "Arial"
REOPEN_GAP = 600
MIN_FLAP = 20

def fmt(n):
    if pd.isna(n): return "N/A"
    if isinstance(n,float) and n==int(n): return f"{int(n):,}"
    if isinstance(n,float): return f"{n:,.1f}"
    return f"{int(n):,}"

def fpct(n):
    return "N/A" if pd.isna(n) else f"{n:.1f}%"

def fdur(s):
    if pd.isna(s): return "N/A"
    if s<60: return f"{s:.0f}s"
    if s<3600: return f"{s/60:.1f}m"
    return f"{s/3600:.1f}h"

def shade(cell,color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def add_tbl(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for ci,h in enumerate(headers):
        c=t.rows[0].cells[ci]; c.text=''; r=c.paragraphs[0].add_run(str(h))
        r.font.name=FONT; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(c,HEADER_BG)
    for ri,row in enumerate(rows):
        bg=ROW_ALT_BG if ri%2==0 else WHITE_BG
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; r=c.paragraphs[0].add_run(str(v))
            r.font.name=FONT; r.font.size=Pt(9); r.font.color.rgb=DARK; shade(c,bg)
    if widths:
        for ri in range(len(t.rows)):
            for ci,w in enumerate(widths):
                t.rows[ri].cells[ci].width=Inches(w)
    doc.add_paragraph()
    return t

def hd(doc,txt,level=1):
    h=doc.add_heading(txt,level=level)
    for r in h.runs: r.font.name=FONT; r.font.color.rgb=PRIMARY if level==1 else SECONDARY if level==2 else DARK

def pa(doc,txt,**kw):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.name=FONT
    r.font.size=kw.get('size',Pt(11)); r.font.color.rgb=kw.get('color',DARK)
    r.italic=kw.get('italic',False); r.bold=kw.get('bold',False)

def kpi_row(doc,kpis):
    t=doc.add_table(rows=2,cols=len(kpis)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ci,(label,val) in enumerate(kpis):
        cv=t.rows[0].cells[ci]; cv.text=''; p=cv.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(val)); r.font.name=FONT; r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=PRIMARY
        shade(cv,"F8FAF9")
        cl=t.rows[1].cells[ci]; cl.text=''; p=cl.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(label)); r.font.name=FONT; r.font.size=Pt(9); r.font.color.rgb=GRAY
        shade(cl,"F8FAF9")
    doc.add_paragraph()

# === ANALYSIS ===
def load_data(csv_path):
    print(f"Loading {csv_path}...")
    df = pd.read_csv(csv_path, low_memory=False)
    df['ts'] = pd.to_datetime(df['timestamp'], unit='ms')
    required = ['timestamp','event','incidentId','conditionName','policyName','durationSeconds','targetName','priority']
    missing = [c for c in required if c not in df.columns]
    if missing:
        print(f"ERROR: Missing columns: {missing}"); sys.exit(1)
    print(f"  Rows: {len(df):,}, Incidents: {df['incidentId'].nunique():,}, Range: {df['ts'].min().date()} to {df['ts'].max().date()}")
    return df

def analyze_all(df):
    opens=df[df['event']=='open']; closes=df[df['event']=='close']
    cdur=pd.to_numeric(closes['durationSeconds'],errors='coerce').dropna()
    evts=df.groupby('incidentId')['event'].apply(list)
    both=sum(1 for e in evts if 'open' in e and 'close' in e)
    oonly=sum(1 for e in evts if 'open' in e and 'close' not in e)
    conly=sum(1 for e in evts if 'close' in e and 'open' not in e)
    days=(df['ts'].max()-df['ts'].min()).days
    he=df['entity.name'].notna().sum() if 'entity.name' in df.columns else 0
    ht=df['targetName'].notna().sum() if 'targetName' in df.columns else 0
    def xsev(n):
        m=re.search(r'SEV(\d)',str(n)); return f'SEV{m.group(1)}' if m else 'Unknown'
    sev_d=opens.copy(); sev_d['sev']=sev_d['policyName'].apply(xsev)
    sev_c=sev_d.groupby('sev').size().sort_values(ascending=False).to_dict()
    S={'rows':len(df),'opens':len(opens),'closes':len(closes),'incidents':df['incidentId'].nunique(),
       'conditions':df['conditionName'].nunique(),'policies':df['policyName'].nunique(),
       'ds':df['ts'].min().date(),'de':df['ts'].max().date(),'days':days,
       'both':both,'oonly':oonly,'conly':conly,
       'ent_pct':he/len(df)*100,'tgt_pct':ht/len(df)*100,
       'pri':df['priority'].value_counts().to_dict(),'sev':sev_c,
       'med_dur':cdur.median() if len(cdur) else None,'mean_dur':cdur.mean() if len(cdur) else None,
       'pct5':((cdur<=300).sum()/len(cdur)*100) if len(cdur) else None,
       'muted':(df['muted']==True).sum()/len(df)*100 if 'muted' in df.columns else None,
       'cond_ids':df['conditionId'].nunique() if 'conditionId' in df.columns else None,
       'pol_ids':df['policyId'].nunique() if 'policyId' in df.columns else None}

    # Conditions
    csev=opens.groupby('conditionName').apply(lambda g:re.search(r'SEV(\d)',g['policyName'].mode().iloc[0]) if len(g['policyName'].mode())>0 else None)
    cnd=opens.groupby('conditionName').agg(oc=('incidentId','count'),ui=('incidentId','nunique'),
        ut=('targetName','nunique'),pl=('policyName','nunique'),
        et=('entity.type',lambda x:x.dropna().mode().iloc[0] if len(x.dropna())>0 else 'NRQL')
    ).sort_values('oc',ascending=False).head(20).reset_index()
    cnd['sev']=cnd['conditionName'].map(lambda c:f"SEV{csev[c].group(1)}" if c in csev.index and csev[c] else 'Unk')
    tc=cnd.iloc[0] if len(cnd)>0 else None
    tp=tc['oc']/len(opens)*100 if tc is not None else 0

    # Policies
    pol=opens.groupby('policyName').agg(oc=('incidentId','count'),uc=('conditionName','nunique'),
        ut=('targetName','nunique')).sort_values('oc',ascending=False).head(15).reset_index()
    cp=df.groupby('conditionName')['policyName'].nunique().reset_index(name='pc')
    dup=cp[cp['pc']>1].sort_values('pc',ascending=False).head(10)

    # Flappiness
    cv=closes.copy(); cv['dur']=pd.to_numeric(cv['durationSeconds'],errors='coerce'); vv=cv[cv['dur'].notna()]
    tot=len(vv)
    fs={'tot':tot,'u1':((vv['dur']<=60).sum()/tot*100) if tot else 0,
        'u5':((vv['dur']<=300).sum()/tot*100) if tot else 0,
        'med':vv['dur'].median() if tot else 0,'mean':vv['dur'].mean() if tot else 0} if tot else {}
    cf=vv.groupby('conditionName').agg(t=('dur','count'),u5=('dur',lambda x:(x<=300).sum()),
        md=('dur','median')).reset_index()
    cf['p5']=cf['u5']/cf['t']*100; cf=cf[cf['t']>=MIN_FLAP].sort_values('p5',ascending=False).head(15)

    # Re-open
    df2=df.copy()
    if 'entity.name' in df2.columns: df2['ee']=df2['entity.name'].fillna(df2['targetName'])
    else: df2['ee']=df2['targetName']
    tcn=opens['conditionName'].value_counts().index[0] if len(opens)>0 else None
    tcp=opens['conditionName'].value_counts().iloc[0]/len(opens)*100 if len(opens)>0 else 0
    exc=tcn if tcp>50 else None
    if exc: df2=df2[df2['conditionName']!=exc]
    cl2=df2[df2['event']=='close'][['ts','conditionName','ee']].sort_values('ts')
    op2=df2[df2['event']=='open'][['ts','conditionName','ee']].sort_values('ts')
    grps=df2.groupby(['conditionName','ee']).size(); grps=grps[grps>=4].reset_index()
    rr=[]
    for _,row in grps.iterrows():
        cn,en=row['conditionName'],row['ee']
        gc=cl2[(cl2['conditionName']==cn)&(cl2['ee']==en)].sort_values('ts')
        go=op2[(op2['conditionName']==cn)&(op2['ee']==en)].sort_values('ts')
        if len(gc)<2 or len(go)<2: continue
        ct,ot=gc['ts'].values,go['ts'].values
        gaps=[((ot[ot>c][0]-c)/np.timedelta64(1,'s')) for c in ct if len(ot[ot>c])>0 and (ot[ot>c][0]-c)/np.timedelta64(1,'s')<3600]
        if len(gaps)>=2:
            rr.append({'cn':cn,'ee':en,'tc':len(gc),'r10':sum(1 for g in gaps if g<REOPEN_GAP),
                        'r30':sum(1 for g in gaps if g<1800),'mg':np.median(gaps)})
    rd=pd.DataFrame(rr)
    if len(rd)>0:
        rd['rate']=rd['r10']/rd['tc']*100; rd=rd.sort_values('r10',ascending=False)
        ra=rd.groupby('cn').agg(ents=('ee','count'),tr=('r10','sum'),tc=('tc','sum'),mg=('mg','median')).reset_index()
        ra['rate']=ra['tr']/ra['tc']*100; ra=ra.sort_values('tr',ascending=False)
    else: ra=pd.DataFrame()

    # Expiration
    cc=closes['closeCause'].value_counts().to_dict()
    vtl=df.groupby('conditionName')['violationTimeLimitSeconds'].first().reset_index()
    vtld=vtl['violationTimeLimitSeconds'].value_counts().sort_index().to_dict()
    cve_t=(df.groupby('conditionName')['closeViolationsOnExpiration'].first()==True).sum()
    cve_f=(df.groupby('conditionName')['closeViolationsOnExpiration'].first()==False).sum()
    cd2=closes.copy(); cd2['dur']=pd.to_numeric(cd2['durationSeconds'],errors='coerce')
    lng=cd2[cd2['dur']>43200]; vlng=cd2[cd2['dur']>86400]
    lrc=pd.DataFrame()
    if len(lng)>0:
        lrc=lng.groupby('conditionName').agg(ct=('incidentId','count'),mx=('dur',lambda x:x.max()/3600),
            av=('dur',lambda x:x.mean()/3600)).sort_values('ct',ascending=False).head(10).reset_index()

    # Entities
    op3=opens.copy()
    if 'entity.name' in op3.columns: op3['ee']=op3['entity.name'].fillna(op3['targetName'])
    else: op3['ee']=op3['targetName']
    eexc=None
    if tp>50: op3=op3[op3['conditionName']!=cnd.iloc[0]['conditionName']]; eexc=cnd.iloc[0]['conditionName']
    ent=op3.groupby('ee').agg(oc=('incidentId','count'),uc=('conditionName','nunique'),
        et=('entity.type',lambda x:x.dropna().mode().iloc[0] if len(x.dropna())>0 else 'NRQL target')
    ).sort_values('oc',ascending=False).head(15).reset_index()

    # Targets
    tn_all=opens.groupby('targetName').agg(oc=('incidentId','count'),uc=('conditionName','nunique'),
        up=('policyName','nunique'),he=('entity.name',lambda x:x.notna().any()) if 'entity.name' in opens.columns else ('targetName',lambda x:False)
    ).sort_values('oc',ascending=False).head(20).reset_index()
    op4=opens.copy()
    texc=None
    if tp>50: op4=op4[op4['conditionName']!=cnd.iloc[0]['conditionName']]; texc=cnd.iloc[0]['conditionName']
    tn_cln=op4.groupby('targetName').agg(oc=('incidentId','count'),uc=('conditionName','nunique'),
        he=('entity.name',lambda x:x.notna().any()) if 'entity.name' in op4.columns else ('targetName',lambda x:False),
        et=('entity.type',lambda x:x.dropna().mode().iloc[0] if len(x.dropna())>0 else 'NRQL FACET')
    ).sort_values('oc',ascending=False).head(15).reset_index()
    at=opens.groupby('targetName').agg(he=('entity.name',lambda x:x.notna().any()) if 'entity.name' in opens.columns else ('targetName',lambda x:False))
    mt=at['he'].sum(); tt=len(at)

    # Entity types
    op5=opens.copy()
    if 'entity.type' in op5.columns: op5['et2']=op5['entity.type'].fillna('(No Entity / NRQL)')
    else: op5['et2']='(No Entity / NRQL)'
    etd=op5.groupby('et2').agg(oc=('incidentId','count'),uc=('conditionName','nunique'),
        ue=('entity.name','nunique') if 'entity.name' in op5.columns else ('targetName','nunique')
    ).sort_values('oc',ascending=False).reset_index()
    etd['pct']=etd['oc']/len(opens)*100

    # Fields
    flds=[]
    for col in df.columns:
        nc=df[col].isna().sum(); nn=len(df)-nc; u=df[col].nunique()
        fi={'f':col,'dt':str(df[col].dtype)[:8],'nn':nn,'nl':nc,'pp':nn/len(df)*100,'u':u}
        tv=df[col].value_counts(dropna=False).head(3)
        fi['tv']=[(str(v)[:40],c) for v,c in tv.items()]
        if pd.api.types.is_numeric_dtype(df[col]):
            vl=df[col].dropna()
            if len(vl): fi['z']=(vl==0).sum(); fi['mn']=vl.min(); fi['mx']=vl.max()
        flds.append(fi)

    return (S,cnd,tc,tp,pol,dup,(cp['pc']>1).sum(),fs,cf,rd.head(15) if len(rd)>0 else pd.DataFrame(),
            ra.head(15) if len(ra)>0 else pd.DataFrame(),exc,
            cc,vtld,cve_t,cve_f,len(lng),len(vlng),lrc,
            ent,eexc,tn_all,tn_cln,texc,mt,tt,etd,flds)

# === REPORT ===
def gen_report(acct,analyst,A,out):
    (S,cnd,tc,tp,pol,dup,ndup,fs,cf,rd,ra,rexc,cc,vtld,cvt,cvf,nl12,nl24,lrc,ent,eexc,
     tna,tnc,texc,mt,tt,etd,flds)=A
    doc=Document()
    st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(11)

    # Cover
    doc.add_paragraph(); doc.add_paragraph()
    h=doc.add_heading('ALERT QUALITY',level=1)
    for r in h.runs: r.font.name=FONT; r.font.size=Pt(28); r.font.color.rgb=PRIMARY
    h=doc.add_heading('MANAGEMENT',level=1)
    for r in h.runs: r.font.name=FONT; r.font.size=Pt(28); r.font.color.rgb=PRIMARY
    pa(doc,"Deep-Dive Analysis & Workshop Preparation",color=SECONDARY,size=Pt(14))
    doc.add_paragraph()
    pa(doc,acct,color=SECONDARY,size=Pt(12))
    pa(doc,f"Analysis Period: {S['ds']} to {S['de']} (~{S['days']} days)",color=GRAY,size=Pt(10))
    doc.add_paragraph()
    pa(doc,f"Prepared by: {analyst}",color=GRAY,size=Pt(10))
    pa(doc,f"Generated: {datetime.now().strftime('%B %Y')}",color=GRAY,size=Pt(10))
    doc.add_paragraph()
    pa(doc,"CONFIDENTIAL",italic=True,bold=True,color=RED,size=Pt(9))
    doc.add_page_break()

    # S1
    hd(doc,"1. Executive Summary")
    pa(doc,f"This report analyzes alert quality for {acct} over {S['days']} days ({S['ds']} to {S['de']}). "
       f"The dataset contains {fmt(S['rows'])} NrAiIncident events covering {fmt(S['incidents'])} unique incidents "
       f"across {fmt(S['conditions'])} conditions and {fmt(S['policies'])} policies.")
    hd(doc,"Data Model",level=2)
    pa(doc,f"Of {fmt(S['incidents'])} incidents: {fmt(S['both'])} have open+close events, "
       f"{fmt(S['oonly'])} open-only, {fmt(S['conly'])} close-only. "
       f"entity.name populated on {fpct(S['ent_pct'])} of rows; targetName on {fpct(S['tgt_pct'])}.")
    kpi_row(doc,[("Incidents",fmt(S['incidents'])),("Opens",fmt(S['opens'])),
                 ("Conditions",fmt(S['conditions'])),("Policies",fmt(S['policies']))])
    if S['med_dur'] is not None:
        kpi_row(doc,[("Median Dur",fdur(S['med_dur'])),("Mean Dur",fdur(S['mean_dur'])),
                     ("% <5min",fpct(S['pct5'])),("Muted",fpct(S['muted']))])
    if tp>50:
        pa(doc,f"Dominant condition: \"{tc['conditionName']}\" = {fpct(tp)} of all opens.",bold=True,color=RED)
    hd(doc,"Severity (from Policy Names)",level=2)
    add_tbl(doc,["Severity","Opens","% Total"],
            [[sv,fmt(ct),fpct(ct/S['opens']*100)] for sv,ct in S['sev'].items()])
    if len(S['pri'])==1:
        pa(doc,f"All events carry priority=\"{list(S['pri'].keys())[0]}\" — no differentiation.",italic=True,color=AMBER)
    doc.add_page_break()

    # S2
    hd(doc,"2. Noisiest Alert Conditions")
    rows=[[str(i+1),str(r['conditionName'])[:55],fmt(r['oc']),fmt(r['ut']),str(r['sev'])]
          for i,(_,r) in enumerate(cnd.iterrows())]
    add_tbl(doc,["#","Condition","Opens","Targets","SEV"],rows,[0.3,3.5,0.8,0.7,0.6])
    doc.add_page_break()

    # S3
    hd(doc,"3. Noisiest Alert Policies")
    rows=[[str(i+1),str(r['policyName'])[:55],fmt(r['oc']),fmt(r['uc'])]
          for i,(_,r) in enumerate(pol.iterrows())]
    add_tbl(doc,["#","Policy","Opens","Conds"],rows,[0.3,4.0,0.8,0.6])
    if len(dup)>0:
        hd(doc,"Condition Replication",level=2)
        pa(doc,f"{ndup} conditions appear in multiple policies:")
        add_tbl(doc,["Condition","# Policies"],
                [[str(r['conditionName'])[:55],str(r['pc'])] for _,r in dup.iterrows()])
    doc.add_page_break()

    # S4
    hd(doc,"4. Flappiness Analysis")
    if fs:
        kpi_row(doc,[("<1min",fpct(fs['u1'])),("<5min",fpct(fs['u5'])),
                     ("Median",fdur(fs['med'])),("Mean",fdur(fs['mean']))])
        hd(doc,"Flappiest Conditions",level=2)
        rows=[[str(r['conditionName'])[:50],fmt(r['t']),fpct(r['p5']),fdur(r['md'])]
              for _,r in cf.iterrows()]
        add_tbl(doc,["Condition","Closed","%<5min","Median"],rows,[3.5,0.7,0.8,0.8])
    doc.add_page_break()

    # S5
    hd(doc,"5. Re-Open Pattern Analysis")
    if rexc: pa(doc,f"Excluded dominant condition: \"{rexc}\"",italic=True,color=GRAY)
    if len(rd)>0:
        hd(doc,"Top Re-Open Offenders",level=2)
        rows=[[str(r['cn'])[:35],str(r['ee'])[:30],fmt(r['tc']),fmt(r['r10']),
               fpct(r['rate']),fdur(r['mg'])] for _,r in rd.iterrows()]
        add_tbl(doc,["Condition","Target","Closes","<10m","Rate","Gap"],rows,[2.2,1.8,0.6,0.5,0.5,0.5])
    if len(ra)>0:
        hd(doc,"Aggregated by Condition",level=2)
        rows=[[str(r['cn'])[:45],fmt(r['ents']),fmt(r['tr']),fmt(r['tc']),
               fpct(r['rate']),fdur(r['mg'])] for _,r in ra.iterrows()]
        add_tbl(doc,["Condition","Entities","Re-opens","Closes","Rate","Gap"],rows,[2.8,0.6,0.7,0.7,0.5,0.7])
    doc.add_page_break()

    # S6
    hd(doc,"6. Expiration & VTL Configuration")
    hd(doc,"VTL Distribution",level=2)
    rows=[]
    for v,c in sorted(vtld.items()):
        h=v/3600; lbl=f"{h/24:.0f}d ({h:.0f}h)" if h>=24 else f"{h:.0f}h" if h>=1 else f"{v}s"
        rows.append([lbl,fmt(c)])
    add_tbl(doc,["VTL","# Conditions"],rows)
    hd(doc,"Close Causes",level=2)
    add_tbl(doc,["Cause","Count"],[[k,fmt(v)] for k,v in cc.items()])
    pa(doc,f"closeViolationsOnExpiration: {cvt} enabled, {cvf} disabled.")
    if len(lrc)>0:
        hd(doc,f"Long-Running (>12h): {fmt(nl12)} incidents",level=2)
        rows=[[str(r['conditionName'])[:50],fmt(r['ct']),f"{r['mx']:.1f}",f"{r['av']:.1f}"]
              for _,r in lrc.iterrows()]
        add_tbl(doc,["Condition","Count","Max(h)","Avg(h)"],rows)
    doc.add_page_break()

    # S7
    hd(doc,"7. Noisiest Entities")
    if eexc: pa(doc,f"Excluding: \"{eexc}\"",italic=True,color=GRAY)
    rows=[[str(i+1),str(r['ee'])[:50],fmt(r['oc']),fmt(r['uc']),str(r['et'])[:20]]
          for i,(_,r) in enumerate(ent.iterrows())]
    add_tbl(doc,["#","Entity/Target","Opens","Conds","Type"],rows,[0.3,3.0,0.7,0.6,1.2])

    # S8
    hd(doc,"8. Noisiest Signal Targets")
    hd(doc,"All Targets",level=2)
    rows=[[str(i+1),str(r['targetName'])[:50],fmt(r['oc']),fmt(r['uc']),"Y" if r['he'] else "N"]
          for i,(_,r) in enumerate(tna.iterrows())]
    add_tbl(doc,["#","Target","Opens","Conds","Entity"],rows,[0.3,3.2,0.7,0.6,0.6])
    if texc: hd(doc,f"Excluding: {texc[:40]}...",level=2)
    else: hd(doc,"Clean View",level=2)
    rows=[[str(i+1),str(r['targetName'])[:50],fmt(r['oc']),fmt(r['uc']),str(r['et'])[:18]]
          for i,(_,r) in enumerate(tnc.iterrows())]
    add_tbl(doc,["#","Target","Opens","Conds","Type"],rows,[0.3,3.2,0.7,0.6,1.0])
    pa(doc,f"Entity mapping: {mt}/{tt} targets ({mt/tt*100:.1f}%) resolve to NR entities.",italic=True,color=GRAY)
    doc.add_page_break()

    # S9
    hd(doc,"9. Noise by Entity Type")
    rows=[[str(r['et2']),fmt(r['oc']),fpct(r['pct']),fmt(r['uc']),fmt(r['ue'])]
          for _,r in etd.iterrows()]
    add_tbl(doc,["Type","Opens","% Total","Conds","Entities"],rows)
    doc.add_page_break()

    # S10
    hd(doc,"10. Prioritized Recommendations")
    pa(doc,"[Use the companion _prompt.md file to generate data-driven recommendations via Claude.]",
       italic=True,color=AMBER)
    doc.add_page_break()

    # Appendix A
    hd(doc,"Appendix A: Workshop Session Guide")
    for title,desc in [("Session 1: Baseline Review (30 min)","Walk through AQM KPIs and this report."),
                        ("Session 2: Top Offender Deep Dives (45 min)","Review top 5 conditions in New Relic. Assign owners."),
                        ("Session 3: Structural Improvements (30 min)","Priority mapping, VTL, condition replication."),
                        ("Session 4: Action Plan (15 min)","Document actions, owners, dates.")]:
        hd(doc,title,level=2); pa(doc,desc)
    doc.add_page_break()

    # Appendix B
    hd(doc,"Appendix B: Methodology")
    pa(doc,f"Data from NrAiIncident for {acct}, {S['ds']} to {S['de']}. "
       f"{fmt(S['rows'])} events: {fmt(S['opens'])} open, {fmt(S['closes'])} close, "
       f"{fmt(S['incidents'])} unique incidents. "
       f"Flappiness from durationSeconds on close events. Re-open analysis: close-to-next-open gaps <10min.")
    doc.add_page_break()

    # Appendix C
    hd(doc,"Appendix C: Field Analysis")
    pa(doc,f"All {len(flds)} fields in the raw export:")
    rows=[]
    for f in flds:
        findings=[]
        if f['pp']==0: findings.append("100% null")
        elif f['pp']<1: findings.append(f"~empty ({f['pp']:.1f}%)")
        if f['u']==1 and f['pp']>50:
            v=f['tv'][0][0] if f['tv'] else '?'; findings.append(f"Const: {v[:25]}")
        if 'z' in f and f['nn']>0 and f['z']/f['nn']>0.8: findings.append(f"{f['z']/f['nn']*100:.0f}% zeros")
        fs2="; ".join(findings) if findings else "Normal"
        rows.append([f['f'],f['dt'],fpct(f['pp']),fmt(f['u']),fs2[:40]])
    add_tbl(doc,["Field","Type","Pop%","Unique","Finding"],rows,[2.0,0.6,0.5,0.6,2.5])

    doc.save(out); print(f"Report: {out}")

# === PROMPT ===
def gen_prompt(acct,A,out):
    (S,cnd,tc,tp,pol,dup,ndup,fs,cf,rd,ra,rexc,cc,vtld,cvt,cvf,nl12,nl24,lrc,ent,eexc,
     tna,tnc,texc,mt,tt,etd,flds)=A
    L=[]
    L.append("# AQM Report — LLM Prompt for Narrative Generation\n")
    L.append("## Instructions\n")
    L.append("You are a Principal Solution Architect at New Relic producing a professional")
    L.append("AQM consulting report for a customer workshop. Write in a consultative,")
    L.append("data-driven tone worth thousands of dollars in consulting fees.\n")
    L.append("**Terminology:** incident (not alert/violation), condition (rule), policy (grouping),")
    L.append("flappiness (short-lived oscillating incidents), re-open rate, signal target / targetName,")
    L.append("entity (cataloged NR entity via entity.name), VTL (violation time limit).")
    L.append("AQM KPIs: Incident Count, Accumulated Minutes, MTTC, % Under 5 Min, % Investigated, MTTI.\n")
    L.append("**Output:** For each section, provide 1-3 interpretive paragraphs. For Section 10,")
    L.append("generate 10 prioritized actions with estimated impact. Reference specific numbers and names.\n")
    L.append(f"## Account: {acct}")
    L.append(f"## Period: {S['ds']} to {S['de']} (~{S['days']} days)\n")

    L.append("## S1: Summary")
    for k,v in S.items(): L.append(f"- {k}: {v}")
    if tc is not None: L.append(f"- Dominant: \"{tc['conditionName']}\" = {fpct(tp)}")

    L.append("\n## S2: Noisiest Conditions")
    for _,r in cnd.head(10).iterrows():
        L.append(f"  [{fmt(r['oc'])}] {r['conditionName']} (targets:{r['ut']}, sev:{r['sev']})")

    L.append("\n## S3: Noisiest Policies")
    for _,r in pol.head(10).iterrows():
        L.append(f"  [{fmt(r['oc'])}] {r['policyName']} (conds:{r['uc']})")
    L.append(f"- {ndup} duplicated conditions. Top: " +
             ", ".join(f"{r['conditionName']}({r['pc']})" for _,r in dup.head(5).iterrows()))

    L.append("\n## S4: Flappiness")
    if fs: L.append(f"- <1min:{fpct(fs['u1'])}, <5min:{fpct(fs['u5'])}, med:{fdur(fs['med'])}, mean:{fdur(fs['mean'])}")
    for _,r in cf.head(10).iterrows():
        L.append(f"  [{fpct(r['p5'])}] {r['conditionName']} (n={fmt(r['t'])}, med={fdur(r['md'])})")

    L.append("\n## S5: Re-Open")
    if rexc: L.append(f"- Excluded: \"{rexc}\"")
    if len(ra)>0:
        for _,r in ra.head(10).iterrows():
            L.append(f"  [{fmt(r['tr'])}] {r['cn']} (ents:{r['ents']}, rate:{fpct(r['rate'])}, gap:{fdur(r['mg'])})")

    L.append("\n## S6: Expiration")
    L.append(f"- Causes: {cc}")
    L.append(f"- VTL: {vtld}")
    L.append(f"- closeOnExpir: {cvt} enabled / {cvf} disabled")
    L.append(f"- >12h: {fmt(nl12)}, >24h: {fmt(nl24)}")
    if len(lrc)>0:
        for _,r in lrc.iterrows(): L.append(f"  [{fmt(r['ct'])}] {r['conditionName']} (max:{r['mx']:.1f}h)")

    L.append("\n## S7-9: Entities/Targets/Types")
    if eexc: L.append(f"- Entity excluded: \"{eexc}\"")
    for _,r in ent.head(10).iterrows(): L.append(f"  [{fmt(r['oc'])}] {r['ee']} ({r['et']})")
    L.append(f"- Targets: {mt}/{tt} mapped ({mt/tt*100:.1f}%)")
    for _,r in etd.iterrows(): L.append(f"  [{fmt(r['oc'])}|{fpct(r['pct'])}] {r['et2']}")

    L.append("\n## Field Anomalies")
    for f in flds:
        an=[]
        if f['pp']==0: an.append("100% null")
        elif f['pp']<1: an.append(f"~empty")
        if f['u']==1 and f['pp']>50: an.append("constant")
        if 'z' in f and f['nn']>0 and f['z']/f['nn']>0.8: an.append(f"{f['z']/f['nn']*100:.0f}% zeros")
        if an: L.append(f"- {f['f']}: {'; '.join(an)}")

    L.append("\n## Generate Section 10: Top 10 Prioritized Actions")
    L.append("Based on ALL data above, produce 10 actions ordered by noise reduction impact.")
    L.append("Each: heading, estimated impact (quantified), 1-2 sentence recommendation.")

    with open(out,'w') as f: f.write('\n'.join(L))
    print(f"Prompt: {out}")

# === MAIN ===
def main():
    p=argparse.ArgumentParser(description="AQM Report Generator")
    p.add_argument("--csv",required=True)
    p.add_argument("--account",required=True)
    p.add_argument("--analyst",default="Jim Hagan, Principal Solution Architect")
    p.add_argument("--output",default=None)
    a=p.parse_args()
    df=load_data(a.csv)
    print("\nAnalyzing...")
    A=analyze_all(df)
    print("Analysis complete.")
    safe=re.sub(r'[^\w\s-]','',a.account).strip().replace(' ','_')
    op=a.output or f"AQM_Analysis_{safe}.docx"
    pp=op.replace('.docx','_prompt.md')
    gen_report(a.account,a.analyst,A,op)
    gen_prompt(a.account,A,pp)
    print(f"\nDone! Feed {pp} to Claude for narrative polish.")

if __name__=="__main__": main()


class AQMReportBuilder:
    def __init__(self, analyzer, analyst_name="Solution Architecture"):
        self.a = analyzer
        self.analyst = analyst_name
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'; style.font.size = Pt(11); style.font.color.rgb = DARK
        for lvl, sz, clr in [('Heading 1', 16, PRIMARY), ('Heading 2', 13, SECONDARY), ('Heading 3', 11.5, DARK)]:
            s = self.doc.styles[lvl]; s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True; s.font.color.rgb = clr

    def _shading(self, cell, hex_c):
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), hex_c); sh.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(sh)

    def _table(self, headers, rows, widths=None, num_cols=None):
        num_cols = num_cols or []
        t = self.doc.add_table(rows=1+len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
        for ci, h in enumerate(headers):
            c = t.rows[0].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(h); r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = 'Arial'
            self._shading(c, HEADER_BG)
        for ri, rd in enumerate(rows):
            bg = ROW_ALT if ri % 2 == 0 else WHITE_HEX
            for ci, v in enumerate(rd):
                c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
                r = p.add_run(str(v)); r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = DARK
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci in num_cols else WD_ALIGN_PARAGRAPH.LEFT
                self._shading(c, bg)
        if widths:
            for row in t.rows:
                for ci, w in enumerate(widths):
                    row.cells[ci].width = Inches(w)
        self.doc.add_paragraph()

    def _para(self, text, italic=False, bold=False, size=11, color=None):
        p = self.doc.add_paragraph(); r = p.add_run(text)
        r.font.size = Pt(size); r.font.name = 'Arial'; r.font.italic = italic; r.font.bold = bold
        r.font.color.rgb = color or DARK

    def _kpi(self, label, value):
        p = self.doc.add_paragraph()
        r = p.add_run(f"{value}"); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = PRIMARY; r.font.name = 'Arial'
        r = p.add_run(f"  {label}"); r.font.size = Pt(10); r.font.color.rgb = MID_GRAY; r.font.name = 'Arial'

    def build(self):
        s = self.a.summary()
        # COVER
        self.doc.add_paragraph()
        p = self.doc.add_paragraph(); r = p.add_run("ALERT QUALITY\nMANAGEMENT")
        r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = PRIMARY; r.font.name = 'Arial'
        self._para("Deep-Dive Analysis & Workshop Preparation", size=14, color=SECONDARY)
        self.doc.add_paragraph()
        self._para(f"{self.a.account_name}", size=12, color=SECONDARY)
        self._para(f"Account(s): {', '.join(str(a) for a in s['account_ids'])}", size=10, color=MID_GRAY)
        self._para(f"Analysis Period: {s['date_min']} to {s['date_max']} (~{s['days']} days)", size=10, color=MID_GRAY)
        self.doc.add_paragraph()
        self._para(f"Prepared by: {self.analyst}", size=10, color=MID_GRAY)
        self._para(f"Generated: {datetime.now().strftime('%B %Y')}", size=10, color=MID_GRAY)
        self.doc.add_paragraph()
        self._para("CONFIDENTIAL", size=9, bold=True, italic=True, color=RED)
        self.doc.add_page_break()

        # 1. EXEC SUMMARY
        self.doc.add_heading("1. Executive Summary", level=1)
        self._para(f"This report analyzes alert quality for {self.a.account_name} over {s['days']} days "
            f"({s['date_min']} to {s['date_max']}). The dataset contains {s['total_rows']:,} NrAiIncident events "
            f"across {s['unique_incidents']:,} unique incidents, {s['unique_conditions']:,} conditions, "
            f"and {s['unique_policies']:,} policies.")
        self.doc.add_heading("Data Model Summary", level=2)
        self._table(["Metric", "Value"], [
            ["Total Event Rows", f"{s['total_rows']:,}"], ["Unique Incidents", f"{s['unique_incidents']:,}"],
            ["Open Events", f"{s['total_opens']:,}"], ["Close Events", f"{s['total_closes']:,}"],
            ["Open+Close Pairs", f"{s['incidents_both']:,}"], ["Open Only", f"{s['incidents_open_only']:,}"],
            ["Close Only", f"{s['incidents_close_only']:,}"], ["Conditions", f"{s['unique_conditions']:,}"],
            ["Policies", f"{s['unique_policies']:,}"], ["Signal Targets", f"{s['unique_targets']:,}"],
            ["Entities (entity.name)", f"{s['unique_entities']:,} ({s['pct_entity']:.1f}% populated)"],
            ["Priority", ", ".join(f"{k}: {v:,}" for k, v in s['priority_dist'].items())],
        ], [3.0, 4.0], [1])

        # 2. CONDITIONS
        self.doc.add_page_break()
        self.doc.add_heading("2. Noisiest Alert Conditions", level=1)
        conds = self.a.noisiest_conditions(20)
        self._para(f"Top condition accounts for {conds.iloc[0]['pct']:.1f}% of all opens.")
        self._table(["#", "Condition Name", "Opens", "Targets", "Policies", "SEV"],
            [[i+1, r['conditionName'][:55], f"{r['opens']:,}", f"{r['targets']:,}", f"{r['policies']}", r['sev']]
             for i, (_, r) in enumerate(conds.iterrows())], [0.35, 3.2, 0.8, 0.7, 0.65, 0.6], [2,3,4])

        # 3. POLICIES
        self.doc.add_page_break()
        self.doc.add_heading("3. Noisiest Alert Policies", level=1)
        pols = self.a.noisiest_policies(15)
        self._table(["#", "Policy Name", "Opens", "Conds", "SEV"],
            [[i+1, r['policyName'][:60], f"{r['opens']:,}", f"{r['conditions']}", r['sev']]
             for i, (_, r) in enumerate(pols.iterrows())], [0.35, 3.8, 0.8, 0.6, 0.65], [2,3])
        self.doc.add_heading("Condition Replication", level=2)
        dupes = self.a.condition_replication()
        if len(dupes) > 0:
            self._table(["Condition Name", "# Policies"],
                [[r['conditionName'][:55], f"{r['policy_count']}"] for _, r in dupes.iterrows()], [5.0, 1.2], [1])

        # 4. FLAPPINESS
        self.doc.add_page_break()
        self.doc.add_heading("4. Flappiness Analysis", level=1)
        fl = self.a.flappiness()
        if fl:
            self._kpi("close in < 1 minute", f"{fl['pct_under_1min']:.1f}%")
            self._kpi("close in < 5 minutes", f"{fl['pct_under_5min']:.1f}%")
            self._kpi("median duration", f"{fl['median_dur_sec']:.0f}s")
            self._kpi("mean duration", f"{fl['mean_dur_min']:.1f} min")
            self.doc.add_heading("Flappiest Conditions", level=2)
            flap = self.a.flappiest_conditions()
            self._table(["Condition", "Closed", "% <5min", "Median(s)", "Mean(s)"],
                [[r['conditionName'][:50], f"{r['total']:,}", f"{r['pct_under5']:.1f}%",
                  f"{r['median_dur']:.0f}", f"{r['mean_dur']:.0f}"] for _, r in flap.iterrows()],
                [3.2, 0.7, 0.8, 0.8, 0.8], [1,2,3,4])

        # 5. RE-OPEN
        self.doc.add_page_break()
        self.doc.add_heading("5. Incident Re-Open Pattern Analysis", level=1)
        self._para("Conditions that close and re-open for the same target within 10 minutes.")
        el, cl = self.a.reopen_analysis()
        if len(el) > 0:
            self.doc.add_heading("Top Re-Open Offenders", level=2)
            self._table(["Condition", "Target", "Closes", "<10min", "Rate", "Med Gap"],
                [[r['conditionName'][:35], str(r['effective_entity'])[:35], f"{r['total_closes']:,}",
                  f"{r['reopens_under_10min']:,}", f"{r['rate']:.0f}%", f"{r['median_gap_sec']:.0f}s"]
                 for _, r in el.iterrows()], [2.2, 2.2, 0.6, 0.6, 0.5, 0.6], [2,3,4,5])
        if len(cl) > 0:
            self.doc.add_heading("Aggregated by Condition", level=2)
            self._table(["Condition", "Entities", "Re-opens", "Closes", "Rate", "Med Gap"],
                [[r['conditionName'][:40], f"{r['entities']}", f"{r['total_reopens']:,}",
                  f"{r['total_closes']:,}", f"{r['rate']:.0f}%", f"{r['median_gap']/60:.1f}m"]
                 for _, r in cl.iterrows()], [2.6, 0.7, 0.9, 0.7, 0.6, 0.8], [1,2,3,4,5])

        # 6. VTL / EXPIRATION
        self.doc.add_page_break()
        self.doc.add_heading("6. Signal Expiration & Violation Time Limit", level=1)
        vtl = self.a.vtl_analysis()
        self._table(["VTL Setting", "Conditions", "% Total"],
            [[v['setting'], f"{v['conditions']}", f"{v['pct']:.1f}%"] for v in vtl], [3, 1.5, 1.5], [1,2])
        self.doc.add_heading("Close Causes", level=2)
        cc = self.a.close_causes()
        self._table(["Cause", "Count", "%"],
            [[k, f"{v:,}", f"{v/s['total_closes']*100:.1f}%"] for k,v in cc.items()], [3, 1.5, 1.5], [1,2])
        lr = self.a.long_running()
        if len(lr) > 0:
            self.doc.add_heading("Long-Running Incidents (>12h)", level=2)
            self._table(["Condition", "Count", "Max(h)", "Avg(h)"],
                [[c[:50], f"{r['count']:.0f}", f"{r['max_hrs']:.1f}", f"{r['avg_hrs']:.1f}"]
                 for c, r in lr.iterrows()], [3.5, 0.8, 1, 1], [1,2,3])
        exp = self.a.expiration_config()
        self.doc.add_heading("Expiration Config", level=2)
        self._table(["Setting", "Enabled", "Disabled"], [
            ["closeViolationsOnExpiration", f"{exp['closeOnExpir_true']}", f"{exp['closeOnExpir_false']}"],
            ["openViolationOnExpiration", f"{exp['openOnExpir_true']}", f"{exp['openOnExpir_false']}"],
            ["expirationDuration (non-zero)", f"{exp['expDur_nonzero']}", f"{exp['expDur_zero']}"],
        ], [3, 1.5, 1.5], [1,2])

        # 7. ENTITIES
        self.doc.add_page_break()
        self.doc.add_heading("7. Noisiest Entities", level=1)
        ent = self.a.noisiest_entities(15, exclude_top=True)
        self._table(["#", "Entity/Target", "Opens", "Conds", "Type"],
            [[i+1, str(n)[:55], f"{r['opens']:,}", f"{r['conditions']}", r['entity_type']]
             for i, (n, r) in enumerate(ent.iterrows())], [0.35, 3.3, 0.7, 0.6, 1.3], [2,3])

        # 8. TARGETS
        self.doc.add_heading("8. Noisiest Signal Targets", level=1)
        ta, _ = self.a.noisiest_targets(20)
        self._table(["#", "Target", "Opens", "Conds", "Mapped", "Type"],
            [[i+1, str(n)[:50], f"{r['opens']:,}", f"{r['conditions']}", r['mapped'], r['entity_type'][:15]]
             for i, (n, r) in enumerate(ta.iterrows())], [0.35, 2.8, 0.7, 0.6, 0.7, 1.1], [2,3])
        te, exc = self.a.noisiest_targets(15, exclude_top=True)
        if exc: self._para(f"Excluding: \"{exc}\"", italic=True, size=9, color=MID_GRAY)
        self._table(["#", "Target", "Opens", "Conds", "Mapped", "Type"],
            [[i+1, str(n)[:50], f"{r['opens']:,}", f"{r['conditions']}", r['mapped'], r['entity_type'][:15]]
             for i, (n, r) in enumerate(te.iterrows())], [0.35, 2.8, 0.7, 0.6, 0.7, 1.1], [2,3])

        # 9. ENTITY TYPE / SEVERITY
        self.doc.add_page_break()
        self.doc.add_heading("9. Noise by Entity Type & Severity", level=1)
        etd = self.a.entity_type_dist()
        self._table(["Entity Type", "Opens", "% Total", "Conditions"],
            [[str(t)[:40], f"{r['opens']:,}", f"{r['pct']:.1f}%", f"{r['conditions']}"]
             for t, r in etd.iterrows()], [3, 1, 1, 1], [1,2,3])
        self.doc.add_heading("Severity (from policy names)", level=2)
        sev = s['sev_dist']
        self._table(["Severity", "Opens", "%"],
            [[k, f"{v:,}", f"{v/s['total_opens']*100:.1f}%"] for k, v in sev.items()], [2, 2, 2], [1,2])
        self.doc.add_heading("Runbook Coverage", level=2)
        rb = self.a.runbook_coverage()
        self._table(["Metric", "Value"], [
            ["With runbook", f"{rb['conditions_with']}/{rb['conditions_total']} ({rb['conditions_with']/rb['conditions_total']*100:.1f}%)"],
            ["Opens with runbook", f"{rb['opens_with']:,}"], ["Opens without", f"{rb['opens_without']:,}"],
        ], [3.5, 3], [1])

        # 10. RECOMMENDATIONS PLACEHOLDER
        self.doc.add_page_break()
        self.doc.add_heading("10. Prioritized Recommendations", level=1)
        self._para("Use the companion Claude Prompt file to generate professional narrative recommendations.", italic=True, color=MID_GRAY)
        self._para("[See: <account>_AQM_Claude_Prompt.md]", bold=True, color=AMBER)

        # APPENDIX A
        self.doc.add_page_break()
        self.doc.add_heading("Appendix A: Workshop Session Guide", level=1)
        self._para("Session 1: Baseline Review (30 min)")
        self._para("Session 2: Top Offender Deep Dives (45 min)")
        self._para("Session 3: Structural Improvements (30 min)")
        self._para("Session 4: Action Plan & Next Steps (15 min)")

        # APPENDIX B
        self.doc.add_heading("Appendix B: Methodology", level=1)
        self._para(f"Data from NrAiIncident for account(s) {', '.join(str(a) for a in s['account_ids'])}. "
            f"{s['total_rows']:,} events: {s['total_opens']:,} open, {s['total_closes']:,} close, "
            f"{s['unique_incidents']:,} unique incidents. Period: {s['date_min']} to {s['date_max']}.")

        # APPENDIX C
        self.doc.add_page_break()
        self.doc.add_heading("Appendix C: Field Analysis", level=1)
        fa = self.a.field_analysis()
        self._table(["Field", "Type", "Pop%", "Unique", "Top Value"],
            [[f['field'], f['dtype'][:10], f"{f['pct_populated']:.1f}%", f"{f['unique']:,}",
              f"{f['top_values'][0][0][:30]} ({f['top_values'][0][1]:,})" if f['top_values'] else "-"]
             for f in fa], [1.8, 0.7, 0.7, 0.7, 2.4], [2,3])
        self.doc.add_heading("Detected Anomalies", level=2)
        for f in fa:
            if f['pct_populated'] == 0: self._para(f"* {f['field']}: 100% null (unused/deprecated)", size=10)
            if f['unique'] == 1 and f['pct_populated'] > 50: self._para(f"* {f['field']}: single value only", size=10)
            if 'zeros' in f and f['populated'] > 0 and f['zeros']/f['populated'] > 0.8:
                self._para(f"* {f['field']}: {f['zeros']/f['populated']*100:.0f}% zeros", size=10)
        m = self.a.muting_analysis()
        if m.get(False, 0) > 0 and m.get(True, 0) == 0:
            self._para("* muted: 100% False - no muting rules active", size=10)
        e = self.a.expiration_config()
        if e['closeOnExpir_true'] < 5:
            self._para(f"* closeViolationsOnExpiration: only {e['closeOnExpir_true']} conditions enabled", size=10)
        return self.doc

    def save(self, path): self.doc.save(path)


class AQMPromptGenerator:
    def __init__(self, analyzer): self.a = analyzer
    def generate(self):
        s = self.a.summary(); fl = self.a.flappiness()
        conds = self.a.noisiest_conditions(20); pols = self.a.noisiest_policies(15)
        flap = self.a.flappiest_conditions(); el, cl = self.a.reopen_analysis()
        vtl = self.a.vtl_analysis(); cc = self.a.close_causes()
        lr = self.a.long_running(); exp = self.a.expiration_config()
        ent = self.a.noisiest_entities(15, exclude_top=True)
        ta, _ = self.a.noisiest_targets(20); te, exc = self.a.noisiest_targets(15, exclude_top=True)
        etd = self.a.entity_type_dist(); rb = self.a.runbook_coverage()
        dupes = self.a.condition_replication()
        L = []
        def W(t=""): L.append(t)
        W("# AQM Report - Claude Narrative Generation Prompt")
        W(); W("## Instructions for Claude"); W()
        W("You are a Principal Solution Architect at New Relic preparing a professional consulting report")
        W("for a customer Alert Quality Management (AQM) workshop. The report should be authoritative,")
        W("data-driven, and actionable."); W()
        W("### Terminology & Style Guide")
        W("- Use 'incident' (not 'alert' or 'violation') for NrAiIncident events")
        W("- Use 'condition' for the alert rule, 'policy' for the grouping container")
        W("- Use 'signal target' or 'targetName' for the FACET value identifying what was measured")
        W("- Use 'entity' only for targets with a populated entity.guid in the NR entity catalog")
        W("- 'Flappiness' = incidents opening/closing rapidly (<5 min), indicating threshold oscillation")
        W("- 'Re-open pattern' = same condition+target closing then opening within minutes")
        W("- VTL = Violation Time Limit (auto-close timer)")
        W("- AQM KPIs: Incident Count, Accumulated Minutes, MTTC, % Under 5 Min, % Investigated, MTTI")
        W("- Reference: https://docs.newrelic.com/docs/tutorial-create-alerts/manage-alert-quality/"); W()
        W("### CRITICAL: Zero Content Bleed-Over")
        W("Use ONLY the data below. Do not reference any other customer, account, or dataset."); W()
        W("### Output Format")
        W("Generate polished prose for each report section:")
        W("1. Executive Summary (2-3 paragraphs)")
        W("2. Key finding for noisiest conditions")
        W("3. Key finding for noisiest policies")
        W("4. Flappiness interpretation")
        W("5. Re-open pattern interpretation")
        W("6. VTL/Expiration interpretation")
        W("7. Entity/Target analysis")
        W("8. Severity & runbook findings")
        W("9. Top 10 Prioritized Recommendations with impact estimates and implementation steps")
        W("10. Additional anomalies worth highlighting"); W()
        W("---"); W()
        W(f"## Raw Data: {self.a.account_name}"); W()
        W(f"### Summary")
        W(f"- Account: {', '.join(str(a) for a in s['account_ids'])}")
        W(f"- Period: {s['date_min']} to {s['date_max']} (~{s['days']} days)")
        W(f"- Rows: {s['total_rows']:,} | Opens: {s['total_opens']:,} | Closes: {s['total_closes']:,}")
        W(f"- Unique Incidents: {s['unique_incidents']:,}")
        W(f"- Pairs: both={s['incidents_both']:,}, open-only={s['incidents_open_only']:,}, close-only={s['incidents_close_only']:,}")
        W(f"- Conditions: {s['unique_conditions']:,} | Policies: {s['unique_policies']:,}")
        W(f"- Targets: {s['unique_targets']:,} | Entities: {s['unique_entities']:,} ({s['pct_entity']:.1f}%)")
        W(f"- Priority: {s['priority_dist']}")
        W(f"- Severity (policy names): {s['sev_dist']}"); W()

        W("### Top 20 Noisiest Conditions")
        for i, (_, r) in enumerate(conds.iterrows()):
            W(f"  {i+1}. [{r['opens']:,} opens, {r['pct']:.1f}%] {r['conditionName']} | targets={r['targets']:,}, policies={r['policies']}, sev={r['sev']}")
        W()
        W("### Top 15 Noisiest Policies")
        for i, (_, r) in enumerate(pols.iterrows()):
            W(f"  {i+1}. [{r['opens']:,}] {r['policyName'][:80]} | conds={r['conditions']}, sev={r['sev']}")
        W()
        W("### Condition Replication")
        for _, r in dupes.iterrows(): W(f"  [{r['policy_count']} policies] {r['conditionName']}")
        W()
        if fl:
            W("### Flappiness")
            W(f"  Closed: {fl['total_closed']:,} | <1min: {fl['under_1min']:,} ({fl['pct_under_1min']:.1f}%) | <5min: {fl['under_5min']:,} ({fl['pct_under_5min']:.1f}%)")
            W(f"  Median: {fl['median_dur_sec']:.0f}s | Mean: {fl['mean_dur_min']:.1f}min"); W()
            W("### Flappiest Conditions")
            for _, r in flap.iterrows(): W(f"  [{r['pct_under5']:.1f}% <5m] {r['conditionName']} (n={r['total']:,}, med={r['median_dur']:.0f}s)")
            W()
        if len(el) > 0:
            W("### Re-Open Patterns (entity level)")
            for _, r in el.iterrows():
                W(f"  [{r['reopens_under_10min']:,} re-opens, {r['rate']:.0f}%] {r['conditionName'][:40]} | {str(r['effective_entity'])[:40]} | closes={r['total_closes']:,}, gap={r['median_gap_sec']:.0f}s")
            W()
        if len(cl) > 0:
            W("### Re-Open Patterns (by condition)")
            for _, r in cl.iterrows():
                W(f"  [{r['total_reopens']:,}, {r['rate']:.0f}%] {r['conditionName'][:50]} | ent={r['entities']}, closes={r['total_closes']:,}, gap={r['median_gap']/60:.1f}m")
            W()
        W("### VTL Distribution")
        for v in vtl: W(f"  {v['setting']}: {v['conditions']} ({v['pct']:.1f}%)")
        W()
        W("### Close Causes")
        for k, v in cc.items(): W(f"  {k}: {v:,}")
        W()
        if len(lr) > 0:
            W("### Long-Running (>12h)")
            for c, r in lr.iterrows(): W(f"  [{r['count']:.0f}] {c[:50]} | max={r['max_hrs']:.1f}h, avg={r['avg_hrs']:.1f}h")
            W()
        W("### Expiration Config")
        W(f"  closeOnExpir: {exp['closeOnExpir_true']} on / {exp['closeOnExpir_false']} off")
        W(f"  openOnExpir: {exp['openOnExpir_true']} on / {exp['openOnExpir_false']} off")
        W(f"  expirationDur: {exp['expDur_nonzero']} non-zero / {exp['expDur_zero']} zero"); W()
        W("### Noisiest Entities (excl top cond)")
        for n, r in ent.iterrows(): W(f"  [{r['opens']:,}] {str(n)[:55]} | {r['entity_type']}, conds={r['conditions']}")
        W()
        W("### Noisiest Targets (all)")
        for n, r in ta.iterrows(): W(f"  [{r['opens']:,}] {str(n)[:55]} | mapped={r['mapped']}, {r['entity_type'][:15]}")
        W()
        W("### Entity Type Distribution")
        for t, r in etd.iterrows(): W(f"  {str(t)[:40]}: {r['opens']:,} ({r['pct']:.1f}%), {r['conditions']} conds")
        W()
        W("### Runbook Coverage")
        W(f"  With: {rb['conditions_with']}/{rb['conditions_total']} ({rb['conditions_with']/rb['conditions_total']*100:.1f}%)")
        W(f"  Opens with: {rb['opens_with']:,} | without: {rb['opens_without']:,}")
        W()
        W("### Muting"); m = self.a.muting_analysis()
        for k, v in m.items(): W(f"  {k}: {v:,}")
        return "\n".join(L)


def main():
    parser = argparse.ArgumentParser(description="Generate AQM Report")
    parser.add_argument("--csv", required=True, help="Path to raw incidents CSV")
    parser.add_argument("--account", required=True, help="Account/customer name")
    parser.add_argument("--analyst", default="Solution Architecture", help="Analyst name")
    parser.add_argument("--prompt-only", action="store_true", help="Only generate Claude prompt")
    parser.add_argument("--output-dir", default=".", help="Output directory")
    args = parser.parse_args()
    safe = re.sub(r'[^\w\s-]', '', args.account).replace(' ', '_')
    out = Path(args.output_dir); out.mkdir(parents=True, exist_ok=True)
    print(f"Loading {args.csv}...")
    a = AQMAnalyzer(args.csv, args.account)
    s = a.summary()
    print(f"  {s['total_rows']:,} rows, {s['unique_incidents']:,} incidents, {s['days']} days")
    print("Generating Claude prompt...")
    prompt = AQMPromptGenerator(a).generate()
    pp = out / f"{safe}_AQM_Claude_Prompt.md"; pp.write_text(prompt)
    print(f"  Saved: {pp}")
    if not args.prompt_only:
        print("Generating report...")
        b = AQMReportBuilder(a, args.analyst); b.build()
        dp = out / f"{safe}_AQM_Analysis.docx"; b.save(str(dp))
        print(f"  Saved: {dp}")
    print("Done!")

if __name__ == "__main__": main()
